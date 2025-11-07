import os
import re
import sys
import json
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from docx.table import _Cell
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_community.document_loaders.parsers import RapidOCRBlobParser
import logging

# Logging setup
logging.basicConfig(
    filename='pre-processing-errors.log',
    level=logging.ERROR,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

def clean_text(text):
    # Cleans raw text by removing extra whitespace, URLs, and line breaks.
    # Inputs: text (str) – unprocessed string extracted from file.
    # Output: (str) – cleaned and normalized text.
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'http\S+|www\S+', '', text)
    text = text.replace('\n', ' ').replace('\t', ' ')
    return text.strip()

def get_paragraphs_from_docx(docx_path):
    # Extracts non-empty paragraphs from a DOCX file.
    # Inputs: docx_path (str) – path to the DOCX file.
    # Output: (list[str]) – list of paragraph texts.
    doc = Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs

def get_tables_from_docx(docx_path):
    # Extracts all tables from a DOCX file as nested lists of cells.
    # Inputs: docx_path (str) – path to the DOCX file.
    # Output: (list[list[list[str]]]) – list of tables, where each table is a list of row lists.
    doc = Document(docx_path)
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)
    return tables

def process_docx_in_chunks(directory_path, chunk_size=10):
    # Reads all DOCX files in a directory, chunks text by paragraph count, cleans each chunk, and saves results to JSONL.
    # Inputs: directory_path (str) – folder path containing DOCX files.
    #         chunk_size (int) – number of paragraphs per chunk.
    # Output: (list[dict]) – list of chunk metadata dictionaries with cleaned text.
    output_file = "../pre-processed-docs.jsonl"
    all_results = []

    with open(output_file, "a", encoding="utf-8") as f:
        for filename in os.listdir(directory_path):
            if not filename.endswith(".docx"):
                continue

            file_path = os.path.join(directory_path, filename)
            print(f"Processing (chunked): {filename}")

            try:
                paragraphs = get_paragraphs_from_docx(file_path)
                for chunk_idx, start_idx in enumerate(range(0, len(paragraphs), chunk_size)):
                    end_idx = min(start_idx + chunk_size, len(paragraphs))
                    chunk_text = ' '.join(paragraphs[start_idx:end_idx])
                    clean = clean_text(chunk_text)

                    page_data = {
                        "text": clean,
                        "metadata": {
                            "filename": filename,
                            "chunk_number": chunk_idx + 1,
                            "paragraph_range": f"{start_idx + 1}-{end_idx}",
                            "character_count": len(clean)
                        }
                    }
                    f.write(json.dumps(page_data, ensure_ascii=False) + "\n")
                    all_results.append(page_data)

            except Exception as e:
                logging.error(f"Failed to process {filename}: {str(e)}")
                continue

    return all_results

def process_pdf_in_chunks(directory_path, chunk_size=10):
    # Extracts text from all PDF files in a directory using PyMuPDF, cleaning and saving by page chunk.
    # Inputs: directory_path (str) – path containing PDF files.
    #         chunk_size (int) – number of pages per chunk.
    # Output: (list[dict]) – list of processed page-level JSON entries.
    output_file = "../pre-processed-docs.jsonl"
    all_results = []

    with open(output_file, "a", encoding="utf-8") as f:
        for filename in os.listdir(directory_path):
            if not filename.endswith(".pdf"):
                continue

            file_path = os.path.join(directory_path, filename)
            print(f"Processing (chunked): {filename}")

            try:
                doc = fitz.open(file_path)
                total_pages = len(doc)
                for start_page in range(0, total_pages, chunk_size):
                    end_page = min(start_page + chunk_size, total_pages)
                    for page_num in range(start_page, end_page):
                        page = doc.load_page(page_num)
                        text = page.get_text("text")
                        clean = clean_text(text)

                        page_data = {
                            "text": clean,
                            "metadata": {
                                "filename": filename,
                                "page_number": page_num + 1,
                                "character_count": len(clean)
                            }
                        }
                        f.write(json.dumps(page_data, ensure_ascii=False) + "\n")
                        all_results.append(page_data)

                doc.close()

            except Exception as e:
                logging.error(f"Failed to process {filename}: {str(e)}")
                continue

    return all_results

def extract_tables_from_docx(directory_path, output_dir="../output-tables"):
    # Extracts all tables from DOCX files in a directory and saves them as CSV files.
    # Inputs: directory_path (str) – folder containing DOCX files.
    #         output_dir (str) – folder where CSVs will be saved.
    # Output: None (writes extracted tables to disk).
    os.makedirs(output_dir, exist_ok=True)
    for filename in os.listdir(directory_path):
        if not filename.endswith(".docx"):
            continue

        file_path = os.path.join(directory_path, filename)
        print(f"Extracting tables (python-docx): {filename}")

        try:
            tables = get_tables_from_docx(file_path)
            table_count = 0

            for idx, table_data in enumerate(tables, start=1):
                try:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    csv_filename = f"{os.path.splitext(filename)[0]}_table{idx}.csv"
                    csv_path = os.path.join(output_dir, csv_filename)
                    df.to_csv(csv_path, index=False, encoding="utf-8")
                    table_count += 1
                except Exception as e:
                    logging.error(f"Failed to extract table {idx} from {filename}: {str(e)}")
                    continue

            if table_count == 0:
                print(f"No tables found in {filename}")
            else:
                print(f"Extracted {table_count} tables from {filename}")

        except Exception as e:
            logging.error(f"Failed to process {filename}: {str(e)}")
            continue

def extract_tables_from_pdf(directory_path, output_dir="../output-tables", chunk_size=10):
    # Extracts tables from PDF files using PyMuPDF page.find_tables() and saves as CSV.
    # Inputs: directory_path (str) – folder containing PDF files.
    #         output_dir (str) – folder where CSV files will be saved.
    #         chunk_size (int) – number of pages processed at a time.
    # Output: None (writes table CSVs to disk).
    os.makedirs(output_dir, exist_ok=True)
    for filename in os.listdir(directory_path):
        if not filename.endswith(".pdf"):
            continue

        file_path = os.path.join(directory_path, filename)
        print(f"Extracting tables (chunked, PyMuPDF only): {filename}")

        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            table_count = 0

            for start_page in range(0, total_pages, chunk_size):
                end_page = min(start_page + chunk_size, total_pages)
                for page_num in range(start_page, end_page):
                    try:
                        page = doc.load_page(page_num)
                        tables = page.find_tables()

                        if tables.tables:
                            for idx, table in enumerate(tables.tables, start=1):
                                df = pd.DataFrame(table.extract())
                                csv_filename = f"{os.path.splitext(filename)[0]}_page{page_num+1}_table{idx}.csv"
                                csv_path = os.path.join(output_dir, csv_filename)
                                df.to_csv(csv_path, index=False, encoding="utf-8")
                                table_count += 1
                    
                    except Exception as e:
                        logging.error(f"Failed to extract tables from {filename} page {page_num+1}: {str(e)}")
                        continue

            doc.close()

            if table_count == 0:
                print(f"No tables found in {filename}")
            else:
                print(f"Extracted {table_count} tables from {filename}")

        except Exception as e:
            logging.error(f"Failed to open/process {filename}: {str(e)}")
            continue


if __name__ == "__main__":
    # Main execution: determines input path and processes accordingly.
    # Inputs: command-line argument (optional path to file or directory).
    # Output: Writes pre-processed text JSONL and table CSVs.
    path = sys.argv[1] if len(sys.argv) > 1 else "./input-docs"
    
    if os.path.isfile(path):
        filename = os.path.basename(path)
        output_file = "../pre-processed-docs.jsonl"
        results = []
        
        if filename.endswith(".pdf"):
            print(f"Processing (chunked): {filename}")
            try:
                doc = fitz.open(path)
                total_pages = len(doc)
                with open(output_file, "w", encoding="utf-8") as f:
                    for page_num in range(total_pages):
                        page = doc.load_page(page_num)
                        text = page.get_text("text")
                        clean = clean_text(text)
                        page_data = {
                            "text": clean,
                            "metadata": {
                                "filename": filename,
                                "page_number": page_num + 1,
                                "character_count": len(clean)
                            }
                        }
                        f.write(json.dumps(page_data, ensure_ascii=False) + "\n")
                        results.append(page_data)
                doc.close()

                print(f"Extracting tables (PyMuPDF only): {filename}")
                doc = fitz.open(path)
                table_count = 0
                output_dir = "../output-tables"
                os.makedirs(output_dir, exist_ok=True)
                for page_num in range(len(doc)):
                    try:
                        page = doc.load_page(page_num)
                        tables = page.find_tables()
                        if tables.tables:
                            for idx, table in enumerate(tables.tables, start=1):
                                df = pd.DataFrame(table.extract())
                                csv_filename = f"{os.path.splitext(filename)[0]}_page{page_num+1}_table{idx}.csv"
                                csv_path = os.path.join(output_dir, csv_filename)
                                df.to_csv(csv_path, index=False, encoding="utf-8")
                                table_count += 1
                    except Exception as e:
                        logging.error(f"Failed to extract tables from {filename} page {page_num+1}: {str(e)}")
                        continue
                doc.close()
                if table_count == 0:
                    print(f"No tables found in {filename}")
                else:
                    print(f"Extracted {table_count} tables from {filename}")
            except Exception as e:
                logging.error(f"Failed to process {filename}: {str(e)}")

        elif filename.endswith(".docx"):
            print(f"Processing (chunked): {filename}")
            try:
                paragraphs = get_paragraphs_from_docx(path)
                chunk_size = 10
                with open(output_file, "w", encoding="utf-8") as f:
                    for chunk_idx, start_idx in enumerate(range(0, len(paragraphs), chunk_size)):
                        end_idx = min(start_idx + chunk_size, len(paragraphs))
                        chunk_text = ' '.join(paragraphs[start_idx:end_idx])
                        clean = clean_text(chunk_text)
                        page_data = {
                            "text": clean,
                            "metadata": {
                                "filename": filename,
                                "chunk_number": chunk_idx + 1,
                                "paragraph_range": f"{start_idx + 1}-{end_idx}",
                                "character_count": len(clean)
                            }
                        }
                        f.write(json.dumps(page_data, ensure_ascii=False) + "\n")
                        results.append(page_data)

                print(f"Extracting tables (python-docx): {filename}")
                tables = get_tables_from_docx(path)
                table_count = 0
                output_dir = "../output-tables"
                os.makedirs(output_dir, exist_ok=True)
                for idx, table_data in enumerate(tables, start=1):
                    try:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        csv_filename = f"{os.path.splitext(filename)[0]}_table{idx}.csv"
                        csv_path = os.path.join(output_dir, csv_filename)
                        df.to_csv(csv_path, index=False, encoding="utf-8")
                        table_count += 1
                    except Exception as e:
                        logging.error(f"Failed to extract table {idx} from {filename}: {str(e)}")
                        continue
                if table_count == 0:
                    print(f"No tables found in {filename}")
                else:
                    print(f"Extracted {table_count} tables from {filename}")
            except Exception as e:
                logging.error(f"Failed to process {filename}: {str(e)}")
        else:
            print(f"Error: {filename} is not a PDF or DOCX file")
    
    elif os.path.isdir(path):
        output_file = "../pre-processed-docs.jsonl"
        with open(output_file, "w", encoding="utf-8") as f:
            pass
        pdf_results = process_pdf_in_chunks(path, chunk_size=10)
        docx_results = process_docx_in_chunks(path, chunk_size=10)
        extract_tables_from_pdf(path, chunk_size=10)
        extract_tables_from_docx(path)
        print(f"Processed {len(pdf_results) + len(docx_results)} documents in total.")
    else:
        print(f"Error: {path} is neither a file nor a directory")
        sys.exit(1)
